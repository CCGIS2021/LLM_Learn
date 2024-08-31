import os
import docx
import pypandoc

def convert_doc_to_docx(input_file, output_file):
    # 使用 python-docx 读取 .doc 文件
    doc = docx.Document(input_file)
    
    # 创建一个新的 .docx 文件
    new_doc = docx.Document()
    
    # 将文本复制到新文档中
    for paragraph in doc.paragraphs:
        new_doc.add_paragraph(paragraph.text)
    
    # 保存为 .docx 文件
    new_doc.save(output_file)

def convert_docx_to_md(input_file, output_file):
    # 使用 pypandoc 进行转换
    output = pypandoc.convert_file(input_file, 'md', format='docx')
    
    # 将转换后的内容写入到 md 文件中
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(output)



# 指定输入和输出文件路径
input_file = ['1.docx','2.docx']
# output_file = 'example.md'

for file_path in input_file:
    print('handling ',file_path)
    if file_path.endswith('.doc'):
        docx_path = file_path.replace('.doc', '.docx')
        convert_doc_to_docx(file_path, docx_path)
        convert_docx_to_md(docx_path, file_path.replace('.doc', '.md'))
        os.remove(docx_path)
    else:
        convert_docx_to_md(file_path, file_path.replace('.docx', '.md'))
